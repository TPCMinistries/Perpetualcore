from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public" / "legal" / "contracts"
OUT.mkdir(parents=True, exist_ok=True)
EFFECTIVE = "August 29, 2026"
COMPANY = "THE PERPETUAL CORE LLC, a New York limited liability company"
INK = RGBColor(23, 23, 27)
PURPLE = RGBColor(81, 70, 199)
MUTED = RGBColor(92, 92, 101)

def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = color

def setup(doc, title, subtitle):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.49)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in [("Heading 1",16,16,8),("Heading 2",13,12,6),("Heading 3",11.5,8,4)]:
        style = doc.styles[name]; style.font.name = "Arial"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = PURPLE
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("PERPETUAL CORE  /  CUSTOMER CONTRACTING"), 8, True, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("Negotiation draft - counsel review recommended"), 8, False, MUTED)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    font(p.add_run("PERPETUAL CORE"), 9, True, PURPLE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    font(p.add_run(title), 24, True, INK)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    font(p.add_run(subtitle), 11, False, MUTED)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    font(p.add_run(f"Form version: {EFFECTIVE}  |  Provider: {COMPANY}"), 9, False, MUTED)
    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(.18); note.paragraph_format.space_after = Pt(14)
    font(note.add_run("IMPORTANT: "), 9, True, PURPLE)
    font(note.add_run("This is a negotiation-ready starting form, not legal advice. Complete all brackets, attach deployment-specific exhibits, and obtain counsel review before signature."), 9, False, INK)

def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        font(p.add_run(bold_prefix), 10.5, True)
        font(p.add_run(text[len(bold_prefix):]), 10.5)
    else: font(p.add_run(text), 10.5)
    return p

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25); p.paragraph_format.space_after = Pt(4)
        font(p.add_run(item), 10.5)

def signatures(doc, labels=("THE PERPETUAL CORE LLC", "CUSTOMER"), page_break=False):
    if page_break:
        doc.add_page_break()
    doc.add_heading("Signatures", level=1)
    table = doc.add_table(rows=5, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    for cell in table._cells: cell.width = Inches(3.2); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [(labels[0], labels[1]), ("By: ______________________________", "By: ______________________________"), ("Name: Lorenzo Daughtry-Chambers", "Name: [NAME]"), ("Title: Authorized Representative", "Title: [TITLE]"), ("Date: ____________________________", "Date: ____________________________")]
    for r, vals in zip(table.rows, rows):
        for cell, value in zip(r.cells, vals):
            cell.text = ""; font(cell.paragraphs[0].add_run(value), 9, vals == rows[0])
    tbl = table._tbl; tblPr = tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"), "nil"); borders.append(el)
    tblPr.append(borders)

def save(doc, slug):
    path = OUT / f"perpetual-core-{slug}.docx"; doc.save(path); return path

def msa():
    d=Document(); setup(d,"Master Services Agreement","Baseline commercial terms for software, implementation, managed services, and advisory work")
    para(d,"This Master Services Agreement (\"Agreement\") is entered as of [EFFECTIVE DATE] between THE PERPETUAL CORE LLC (\"Provider\") and [CUSTOMER LEGAL NAME] (\"Customer\"). Each is a \"Party\" and together the \"Parties.\"")
    sections=[
      ("1. Agreement structure",["This Agreement governs each order form or statement of work (each, an \"Order\"). An Order identifies the services, fees, term, deliverables, assumptions, responsibilities, and acceptance criteria. If terms conflict, the following order controls: signed amendment; Data Processing Addendum or Business Associate Addendum for its subject; Order; this Agreement; then incorporated online policies."]),
      ("2. Services and change control",["Provider will perform the Services with commercially reasonable care and skill. Timelines depend on Customer providing timely access, decisions, content, personnel, and approvals. Either Party may propose a change. No change to scope, fees, schedule, authority, regulated-data handling, or acceptance criteria is binding until documented in a signed change order."]),
      ("3. Customer responsibilities",["Customer will designate an authorized owner; provide lawful access and accurate instructions; obtain necessary rights, notices, consents, and licenses; maintain systems of record and backups unless assigned to Provider; and review AI-assisted or automated output before consequential use. Customer will not submit regulated data unless the applicable Order and written addendum expressly authorize it."]),
      ("4. Fees, expenses, and taxes",["Customer will pay the fees and approved expenses stated in each Order. Unless stated otherwise, invoices are due within 15 days, deposits are non-refundable once work begins, and undisputed late amounts may accrue interest at the lesser of 1.0% per month or the maximum lawful rate. Customer is responsible for applicable sales, use, and similar taxes other than taxes on Provider's net income. Customer must dispute an invoice in writing within 10 days and pay undisputed amounts when due."]),
      ("5. Acceptance",["Deliverables are accepted when Customer confirms acceptance, uses them in production, or does not provide a specific written rejection tied to agreed acceptance criteria within 7 business days after delivery. Provider will use commercially reasonable efforts to correct a timely, valid rejection. New preferences or out-of-scope requirements require change control."]),
      ("6. Intellectual property",["Customer retains Customer Data. Provider retains its platform, software, methods, models, workflows, templates, know-how, pre-existing materials, and general improvements (\"Provider Materials\"). Upon full payment, Customer receives the ownership or license expressly stated in the Order for custom deliverables. To the extent Provider Materials are embedded in a deliverable, Provider grants Customer a non-exclusive, non-transferable (except with a permitted assignment), perpetual license to use them internally with that deliverable. Provider may use generalized skills and know-how that do not disclose Customer Confidential Information."]),
      ("7. Data, security, and AI",["The Privacy Policy applies to personal information. If Provider processes personal data on Customer's behalf, the Parties will execute the applicable Data Processing Addendum. Provider is not a HIPAA business associate unless an authorized Business Associate Addendum is signed for the applicable service. AI output is probabilistic and may be inaccurate, incomplete, biased, outdated, or non-original. Customer remains responsible for qualified review, professional obligations, filings, official decisions, and use of output."]),
      ("8. Confidentiality",["A receiving Party will use the disclosing Party's non-public information only for the relationship, protect it using at least reasonable care, and disclose it only to personnel and providers with a need to know and confidentiality duties. Exclusions apply to information that is public without breach, already known without duty, independently developed, or rightfully received from another source. Required legal disclosure is permitted after notice where lawful. Trade secrets remain protected while they qualify as trade secrets; other obligations last five years after disclosure."]),
      ("9. Warranties and disclaimers",["Each Party warrants it has authority to enter this Agreement. Provider warrants professional services will be performed in a professional and workmanlike manner. Customer's exclusive remedy for breach is re-performance if promptly requested. EXCEPT AS EXPRESSLY STATED, SERVICES AND OUTPUTS ARE PROVIDED \"AS IS\" AND \"AS AVAILABLE,\" AND IMPLIED WARRANTIES ARE DISCLAIMED TO THE MAXIMUM EXTENT PERMITTED BY LAW. Provider does not guarantee funding, awards, revenue, compliance, employment, placement, clinical, legal, financial, or other business outcomes."]),
      ("10. Indemnification",["Each Party will defend the other against third-party claims arising from its gross negligence, willful misconduct, or violation of law. Customer will also defend claims arising from Customer Data, Customer instructions, or Customer's unauthorized or prohibited use. Provider will defend claims that the unmodified paid Service, as provided, directly infringes a U.S. patent, copyright, or trademark, and may modify, replace, obtain rights, or terminate the affected Service with a prorated refund. The protected Party must promptly notify, reasonably cooperate, and allow the indemnifying Party to control the defense, subject to consent for settlements admitting fault or imposing non-monetary obligations."]),
      ("11. Limitation of liability",["TO THE MAXIMUM EXTENT PERMITTED BY LAW, NEITHER PARTY IS LIABLE FOR INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, PUNITIVE, OR CONSEQUENTIAL DAMAGES OR LOST PROFITS, REVENUE, DATA, OR GOODWILL. EXCEPT FOR PAYMENT OBLIGATIONS, INDEMNIFICATION, MISUSE OF INTELLECTUAL PROPERTY, BREACH OF CONFIDENTIALITY, GROSS NEGLIGENCE, WILLFUL MISCONDUCT, OR LIABILITY THAT CANNOT LAWFULLY BE LIMITED, EACH PARTY'S AGGREGATE LIABILITY WILL NOT EXCEED FEES PAID OR PAYABLE UNDER THE AFFECTED ORDER DURING THE 12 MONTHS BEFORE THE EVENT GIVING RISE TO LIABILITY."]),
      ("12. Term and termination",["This Agreement continues until terminated. Either Party may terminate for uncured material breach after 10 days' notice for nonpayment or 30 days' notice for other breach. An Order may state additional termination rights. On termination, Customer will pay amounts earned and non-cancellable commitments through the effective date. Each Party will return or destroy Confidential Information on request, subject to legal retention and backup cycles. Customer must export needed data before access ends unless the Order provides a transition period."]),
      ("13. General",["The Parties are independent contractors. Neither may bind the other. Neither Party is liable for delay caused by events beyond reasonable control, excluding payment obligations. Neither may assign this Agreement except to an affiliate or in connection with a merger, reorganization, or sale of substantially all relevant assets, provided the assignee assumes the obligations. New York law governs; exclusive venue is in state or federal courts in New York County, New York. Before filing, executives will attempt good-faith resolution for 15 days. Notices must be written to the Order contacts. Electronic signatures and counterparts are effective. This Agreement and its Orders are the entire agreement for their subject matter."])
    ]
    for h, ps in sections: d.add_heading(h,1); [para(d,x) for x in ps]
    signatures(d, page_break=True); return save(d,"master-services-agreement")

def sow():
    d=Document(); setup(d,"Statement of Work Template","Use with the Perpetual Core Master Services Agreement")
    fields=[("SOW number","[PC-YYYY-###]"),("Customer","[LEGAL NAME]"),("Effective date","[DATE]"),("Term","[START] through [END]"),("Offer","[Operating System Map / Workflow Proof Sprint / Managed Operating Lane / Other]"),("Customer owner","[NAME / TITLE]"),("Provider owner","[NAME / TITLE]")]
    t=d.add_table(rows=1,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style="Table Grid"
    for label,value in fields:
        cells=t.add_row().cells; font(cells[0].paragraphs[0].add_run(label),9.5,True); font(cells[1].paragraphs[0].add_run(value),9.5)
    t._element.remove(t.rows[0]._element)
    blocks=[("1. Business objective","[Describe the operating problem, affected people, current baseline, and why the work matters now.]"),("2. Scope and deliverables","List each deliverable with its format, owner, due date, dependencies, and what is expressly out of scope."),("3. Success measures and evidence","Define baseline, target, measurement method, source of truth, evaluation date, and the person authorized to verify results. Do not treat configured software or activity counts as business outcomes."),("4. Authority and approval gates","Identify what may be observed, drafted, or automated; what requires Customer review; who may send, submit, publish, approve, pay, sign, or change an official record; and the escalation path."),("5. Data and system boundaries","List systems of record, integrations, data categories, hosting, subprocessors, retention, export, deletion, and prohibited data. State whether a DPA or BAA applies. No PHI is authorized unless the BAA and technical environment are approved."),("6. Project plan","Discovery / design / build / test / launch / training / outcome review. Add dates and owners for each milestone."),("7. Customer responsibilities and assumptions","List required access, stakeholders, decisions, content, technical resources, procurement, and review turnaround. Record assumptions that affect price or schedule."),("8. Acceptance","For each deliverable, define objective acceptance criteria, review period (default 7 business days), correction process, and the authorized acceptor."),("9. Fees and payment schedule","Total: $[AMOUNT]. Deposit: $[AMOUNT] due on signature. Milestones or recurring fees: [DETAILS]. Approved expenses: [DETAILS]. Taxes: [DETAILS]. Work begins after deposit and required access are received."),("10. Change control","Changes to scope, schedule, fee, authority, data handling, or acceptance criteria require a written change order signed by both Parties."),("11. Support and transition","Define launch support, response targets, maintenance exclusions, documentation, training, export, handoff, and any optional managed lane."),("12. Incorporated terms","This SOW is governed by the Master Services Agreement dated [DATE], plus any signed DPA, BAA, security exhibit, and change orders. If there is a conflict, the MSA's order-of-precedence clause applies.")]
    for h,b in blocks:
        if h.startswith("9. "):
            d.add_page_break()
        d.add_heading(h,1); para(d,b)
    signatures(d, page_break=True); return save(d,"statement-of-work-template")

def nda():
    d=Document(); setup(d,"Mutual Nondisclosure Agreement","Two-way confidentiality for evaluation, sales, partnership, and implementation discussions")
    para(d,"This Mutual Nondisclosure Agreement (\"NDA\") is entered as of [EFFECTIVE DATE] by and between THE PERPETUAL CORE LLC and [COUNTERPARTY LEGAL NAME].")
    clauses=[("1. Purpose","The Parties may disclose information to evaluate or pursue [DESCRIBE PURPOSE] (the \"Purpose\")."),("2. Confidential Information","Confidential Information means non-public business, technical, financial, product, security, customer, personnel, and strategic information disclosed in any form that is marked confidential or reasonably should be understood as confidential. It includes notes and analyses derived from such information."),("3. Exclusions","Confidential Information does not include information the receiving Party can document was public without breach, already known without duty, independently developed without use, or rightfully received from a third party without duty."),("4. Protection and use","The receiving Party will use Confidential Information only for the Purpose, protect it with at least reasonable care, and disclose it only to personnel, affiliates, financing sources, and professional advisers who need to know and are bound by confidentiality duties at least as protective."),("5. Required disclosure","A receiving Party may disclose information required by law after prompt notice where legally permitted and reasonable assistance in seeking protective treatment."),("6. Return or destruction","On written request, the receiving Party will return or destroy Confidential Information, except archival copies required by law or retained in routine backups subject to continuing protection."),("7. No license or obligation","No intellectual-property license is granted except the limited right to evaluate information for the Purpose. Neither Party is required to proceed with a transaction, and information is provided without warranty for evaluation."),("8. Term","This NDA begins on the Effective Date. Confidentiality duties continue for five years after each disclosure; trade secrets remain protected while they qualify as trade secrets under applicable law."),("9. Remedies","Unauthorized use or disclosure may cause irreparable harm for which monetary relief is inadequate. A Party may seek appropriate equitable relief in addition to other remedies, without waiving defenses."),("10. General","New York law governs. Exclusive venue lies in state or federal courts in New York County, New York. This NDA is the entire agreement for confidentiality concerning the Purpose and may be amended only in a signed writing. Electronic signatures and counterparts are effective.")]
    for h,b in clauses: d.add_heading(h,1); para(d,b)
    signatures(d,("THE PERPETUAL CORE LLC","COUNTERPARTY")); return save(d,"mutual-nda")

def dpa():
    d=Document(); setup(d,"Data Processing Addendum","Controller-processor and service-provider terms for customer personal data")
    para(d,"This Data Processing Addendum (\"DPA\") forms part of the agreement between THE PERPETUAL CORE LLC (\"Provider\") and [CUSTOMER] (\"Customer\") and applies only to Provider's processing of Customer Personal Data on Customer's behalf.")
    clauses=[("1. Roles and instructions","Customer is the controller or business; Provider is the processor or service provider, except where law assigns another role. Provider will process Customer Personal Data only to provide the contracted services, follow documented lawful instructions, secure the service, and meet legal obligations."),("2. Processing details","Subject matter: [SERVICE]. Duration: [TERM plus deletion/return period]. Nature and purpose: [DESCRIBE]. Data subjects: [DESCRIBE]. Data categories: [DESCRIBE]. Sensitive data: [NONE unless expressly listed]. Customer will not provide data outside these details without a signed update."),("3. Confidentiality and security","Provider will ensure authorized personnel are bound by confidentiality and maintain safeguards appropriate to the risk, including access control, authentication, transmission protection, supported encryption at rest, logging, vulnerability management, backup/recovery practices, and incident procedures as applicable to the deployment. The security exhibit, if any, controls specific commitments."),("4. Subprocessors","Customer authorizes the subprocessors listed in the applicable order or subprocessor notice. Provider will impose data-protection obligations appropriate to the services and remains responsible for their performance to the extent required by law. Customer may object to a new material subprocessor on reasonable data-protection grounds within [10] days of notice; the Parties will seek a reasonable solution."),("5. Rights requests","Taking into account the nature of processing, Provider will reasonably assist Customer with verified data-subject requests where Customer cannot fulfill them through the service. Provider will not independently respond except as instructed or legally required."),("6. Security incidents","Provider will notify Customer without undue delay after confirming a Security Incident affecting Customer Personal Data, provide available information reasonably needed for Customer's obligations, and take reasonable containment and remediation steps. Notice is not an admission of fault. Unsuccessful attempts and events that do not compromise Customer Personal Data are not Security Incidents."),("7. Assessments and assistance","On reasonable request and subject to confidentiality, Provider will provide information reasonably necessary to demonstrate compliance and assist with risk assessments, regulator consultations, and breach obligations in proportion to the processing. Additional or on-site audit work may be subject to reasonable fees unless a confirmed material breach caused the request."),("8. International transfers","The Parties will implement a lawful transfer mechanism where required. Customer is responsible for identifying residency or localization requirements before processing begins."),("9. Return and deletion","At the end of services and on Customer's written choice, Provider will return or delete Customer Personal Data within the period stated in the order, except data required by law or retained in protected backups until overwritten in the ordinary cycle."),("10. U.S. state privacy terms","Where applicable, Provider will not sell or share Customer Personal Data for cross-context behavioral advertising; retain, use, or disclose it outside the business purposes in the agreement; or combine it with personal data from unrelated sources except as permitted by law. Provider will notify Customer if it determines it can no longer meet applicable restrictions."),("11. Priority and liability","This DPA controls for its subject if it conflicts with the agreement. Liability under this DPA is subject to the agreement's limitations unless applicable law prohibits that limitation.")]
    for h,b in clauses: d.add_heading(h,1); para(d,b)
    signatures(d); return save(d,"data-processing-addendum")

def baa():
    d=Document(); setup(d,"Business Associate Addendum","HIPAA addendum - effective only when signed for an approved service and environment")
    para(d,"THIS ADDENDUM DOES NOT AUTHORIZE PHI BY ITSELF. It becomes effective only when signed by both Parties, attached to an agreement identifying the approved service, and after Provider confirms the technical environment and subprocessors authorized for PHI.")
    para(d,"This Business Associate Addendum (\"BAA\") is entered as of [EFFECTIVE DATE] by [COVERED ENTITY / BUSINESS ASSOCIATE] (\"Covered Entity\") and THE PERPETUAL CORE LLC (\"Business Associate\"). Capitalized terms not defined have the meanings in HIPAA.")
    clauses=[("1. Permitted uses and disclosures","Business Associate may use or disclose PHI only to perform the services described in [ORDER/SOW], as permitted by this BAA, or as Required by Law. Business Associate may use PHI for proper management and administration or legal responsibilities only where HIPAA permits and required assurances are obtained."),("2. Safeguards","Business Associate will not use or disclose PHI other than as permitted; will implement appropriate administrative, physical, and technical safeguards; and will comply with applicable Security Rule requirements for electronic PHI."),("3. Reporting","Business Associate will report to Covered Entity any use or disclosure not permitted by this BAA, including a Breach of Unsecured PHI, and any Security Incident, without unreasonable delay and within the timing stated in the attached security exhibit. The Parties may document routine unsuccessful Security Incidents in aggregate."),("4. Subcontractors","Business Associate will ensure that subcontractors that create, receive, maintain, or transmit PHI agree in writing to the same restrictions, conditions, and applicable safeguards."),("5. Individual rights assistance","As applicable to the services, Business Associate will make PHI available to Covered Entity to support access, amendment, and accounting-of-disclosures obligations, in the time and manner reasonably requested. Business Associate will forward requests received directly unless legally required to respond."),("6. Covered Entity duties delegated","To the extent Business Associate performs an obligation of Covered Entity under the Privacy Rule, Business Associate will comply with the requirements that apply to Covered Entity in performing that obligation."),("7. HHS access","Business Associate will make internal practices, books, and records relating to PHI available to the Secretary of HHS for determining compliance."),("8. Minimum necessary and restrictions","Business Associate will limit uses, disclosures, and requests to the minimum necessary where required and comply with restrictions and confidential-communication requests communicated by Covered Entity to the extent applicable to the services."),("9. Term and termination","This BAA continues while Business Associate maintains PHI. Covered Entity may terminate for Business Associate's material violation if cure is not possible or is not completed within the agreement's cure period. On termination, Business Associate will return or destroy PHI if feasible; if infeasible, it will extend protections and limit further use and disclosure to the reason return or destruction is infeasible."),("10. Covered Entity responsibilities","Covered Entity will not request use or disclosure that would violate HIPAA if performed by Covered Entity; will provide applicable notices, restrictions, and permissions; will identify authorized users and data; and will not transmit PHI until Provider confirms the approved environment in writing."),("11. Interpretation and priority","This BAA will be interpreted to permit HIPAA compliance. It controls over conflicting agreement terms for PHI. Other agreement terms, including fees and liability, remain effective to the extent permitted by HIPAA.")]
    for h,b in clauses: d.add_heading(h,1); para(d,b)
    d.add_heading("Required deployment exhibit",1); bullets(d,["Approved service and environment: [REQUIRED]","Permitted PHI and data subjects: [REQUIRED]","Authorized subprocessors and BAAs: [REQUIRED]","Security controls and incident notice timing: [REQUIRED]","Return, export, and deletion procedure: [REQUIRED]","Covered Entity privacy/security contacts: [REQUIRED]"])
    signatures(d,("THE PERPETUAL CORE LLC","COVERED ENTITY")); return save(d,"business-associate-addendum")

if __name__ == "__main__":
    paths=[msa(),sow(),nda(),dpa(),baa()]
    print("\n".join(str(p) for p in paths))
